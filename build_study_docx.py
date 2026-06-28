#!/usr/bin/env python3
"""Generate the vsh interview study guide as a .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------------------------------------------------------------- base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.08

# Color palette
DARK = RGBColor(0x1F, 0x2A, 0x44)
ACCENT = RGBColor(0x1B, 0x5E, 0x20)   # green
CODECLR = RGBColor(0x14, 0x14, 0x14)
GREY = RGBColor(0x55, 0x55, 0x55)


def shade(paragraph, fill):
    """Apply a background shade to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def border(paragraph, color="CCCCCC", size="6", sides=("left",)):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "8")
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = DARK
    r.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = ACCENT
    r.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(10)
    return p


def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = DARK
    r.font.size = Pt(12)
    return p


def para(text="", bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    add_runs(p, text)
    for run in p.runs:
        run.bold = bold or run.bold
        run.italic = italic or run.italic
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    return p


def add_runs(paragraph, text):
    """Add text, interpreting **bold** and `code` inline markup."""
    import re
    tokens = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB0, 0x2A, 0x37)
        else:
            paragraph.add_run(tok)


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    add_runs(p, text)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    add_runs(p, text)
    return p


def code(text):
    """Monospace, shaded code block. Each line is its own paragraph."""
    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = CODECLR
        shade(p, "F2F3F5")
        border(p, color="D9DBDE", sides=("left",))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def callout(label, text):
    """Interview-answer style box."""
    p = doc.add_paragraph()
    shade(p, "E8F1E9")
    border(p, color="1B5E20", size="12", sides=("left",))
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10.5)
    add_runs(p, text)
    for run in p.runs[1:]:
        run.font.size = Pt(10.5)
        run.italic = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            add_runs(cells[i].paragraphs[0], str(val))
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def hr():
    p = doc.add_paragraph()
    border(p, color="999999", size="6", sides=("bottom",))
    p.paragraph_format.space_after = Pt(2)


def pagebreak():
    doc.add_page_break()


# ============================================================ TITLE PAGE
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("vsh — Vanguard Shell")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = DARK
title.paragraph_format.space_before = Pt(120)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Complete Interview Study Guide")
r.font.size = Pt(18)
r.font.color.rgb = ACCENT

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("A from-scratch Linux shell in C — every file, every concept, explained")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(40)
r = meta.add_run("Jacob Antony Jeejo\nAirspan India — Software Development internship prep\nInterview: 25 June 2026, Bengaluru")
r.font.size = Pt(12)
r.font.color.rgb = DARK

stat = doc.add_paragraph()
stat.alignment = WD_ALIGN_PARAGRAPH.CENTER
stat.paragraph_format.space_before = Pt(30)
r = stat.add_run("~9,600 lines of C (src + headers)  •  ~1,400 lines of tests (215 unit tests)\nZero third-party libraries  •  100% C standard library + POSIX")
r.font.size = Pt(10.5)
r.italic = True
r.font.color.rgb = GREY

pagebreak()

# ============================================================ HOW TO USE
h1("How to Use This Guide")
para("You built vsh with AI assistance, so you have never read it end-to-end yourself. This guide fixes that. It assumes you know **basic C** (variables, loops, functions, simple pointers) and teaches you everything else from the ground up so you can explain the project out loud, confidently, to an Airspan interviewer.")
para("Read it in order — it is built like a staircase:")
numbered("**The 60-second pitch** — memorise this first so you always have an answer ready.")
numbered("**Shell concepts** — what a shell actually is and the OS machinery it sits on (fork, exec, pipes, signals). This is the vocabulary the rest of the guide uses.")
numbered("**Resume terms decoded** — every buzzword on your resume (POSIX, recursive-descent parser, arena allocator, AddressSanitizer…) explained in plain English, with the exact sentence to say.")
numbered("**Recursive-descent parsing, in depth** — the interviewer's most likely deep-dive topic, including its drawbacks (they will ask).")
numbered("**Systems-C refresher** — the non-basic C the project leans on (file descriptors, va_list, function pointers).")
numbered("**Architecture + file-by-file walkthrough** — what every file in the repo does.")
numbered("**The 'new' / showcase features** — calc, httpfetch, sysinfo, the git prompt — the things that make vsh stand out.")
numbered("**Known bugs & honest resume audit** — what is real, what is a stretch, and how to talk about both honestly.")
numbered("**Q&A bank + the Airspan/embedded connection + the 4-minute pitch.**")

callout("Time-saver:", "With only days left, read sections 1–4 thoroughly (concepts + parser + resume terms). Skim 6 (file walkthrough) for the modules you are most likely to be quizzed on: arena, parser, executor, pipeline, job_control. Memorise sections 1 and 11.")

pagebreak()

# ============================================================ 1. 60-SECOND PITCH
h1("1.  The 60-Second Pitch")
para("Say this near the start of the interview when they ask \"tell me about a project.\"")
callout("Pitch:",
        "\"vsh is a POSIX-style Unix shell I wrote from scratch in C — about 9,600 lines, with zero third-party "
        "libraries, not even readline. It has four stages: a custom raw-mode line editor reads keystrokes, a "
        "hand-written lexer tokenises the input, a recursive-descent parser builds an abstract syntax tree, and an "
        "executor walks that tree — forking processes, wiring up pipes, and running builtins in-process. The part I'm "
        "proudest of is the memory model: every token and AST node is allocated from an arena allocator, so after each "
        "command one O(1) arena_reset() frees the whole parse tree — there's no way to leak parser memory. It also has "
        "full Unix job control — Ctrl+Z, fg, bg, background jobs — built on process groups and a SIGCHLD handler. It's "
        "compiled under -Wall -Wextra -Werror with AddressSanitizer, and verified by a 215-test suite.\"")
para("Then stop and let them pick what to drill into. Every claim above has a dedicated section below.")

hr()

# ============================================================ 2. SHELL CONCEPTS
h1("2.  Shell Concepts — What a Shell Actually Is")

h2("2.1  What is a shell?")
para("A **shell** is the program that sits between you and the operating system. When you type `ls -la` in a terminal, something has to read that text, find the `ls` program on disk, run it, and show you its output. That \"something\" is the shell.")
para("Shells you already know: **bash** (default on most Linux), **zsh** (macOS default), **fish**. vsh is the same kind of program, written by you.")
para("A shell is fundamentally a **REPL** — a Read–Eval–Print Loop:")
code("""while (shell is running) {
    print a prompt          (e.g. "vsh> ")
    READ a line of input    (vsh_readline.c)
    EVALUate it             (lexer -> parser -> executor)
    PRINT any output        (the program's output goes to the terminal)
}""")

h2("2.2  The four stages of running one command")
para("Everything vsh does maps onto four steps. For `ls -la | grep .txt`:")
table(
    ["Stage", "File", "What it does", "Output"],
    [
        ["1. Read", "vsh_readline.c", "Reads your keystrokes in raw mode", '`"ls -la | grep .txt"`'],
        ["2. Lex", "lexer.c", "Splits the string into tokens", "`[ls] [-la] [|] [grep] [.txt]`"],
        ["3. Parse", "parser.c", "Finds the structure / precedence", "an AST (a tree)"],
        ["4. Execute", "executor.c + pipeline.c", "Forks processes, wires pipes, waits", "ls and grep run"],
    ],
    widths=[1.0, 1.5, 2.7, 1.6],
)

h2("2.3  Processes: fork(), exec(), wait()")
para("This is the single most important idea in the whole project. When you type `ls`, the shell does **not** call a function named `ls`. It creates a brand-new **process** to run the `ls` program, while the shell itself stays alive.")
para("**fork()** makes an exact copy of the current process. After it returns, *two* processes are running the same code:")
code("""pid_t pid = fork();
if (pid == 0) {
    // CHILD: fork() returned 0 here
} else if (pid > 0) {
    // PARENT: fork() returned the child's PID here
} else {
    // fork() failed (pid == -1)
}""")
para("**execvp()** replaces the current process image with a different program. The code, data, everything is swapped out for (say) `ls`. If exec succeeds, the line after it never runs:")
code("""char *args[] = { "ls", "-la", NULL };  // argv, NULL-terminated
execvp("ls", args);
perror("exec failed");   // only reached if exec FAILED""")
para("**Together**, fork + exec is how you run a program without killing the shell:")
code("""Shell (PID 100) calls fork()
   |
   +--> Parent (PID 100): waitpid(101) -- blocks until child finishes
   |
   +--> Child  (PID 101): execvp("ls", args)
                          -> becomes ls, prints files, exits
                          -> parent's waitpid() returns
                          -> shell prints the next prompt""")
callout("Interview line:",
        "\"fork() duplicates the process; exec() replaces it with a new program; waitpid() lets the parent block "
        "until the child finishes. The shell forks a child, execs the target program in the child, and waits in the "
        "parent — so the shell survives to run the next command.\"")

h2("2.4  File descriptors")
para("In Linux, every open I/O channel is a small integer called a **file descriptor (fd)**:")
table(
    ["fd", "Name", "Default meaning"],
    [["0", "stdin", "keyboard input"], ["1", "stdout", "terminal output"],
     ["2", "stderr", "error output"], ["3, 4, …", "(yours)", "files/sockets you open"]],
    widths=[0.8, 1.2, 3.5],
)
para("`open(\"notes.txt\", O_RDONLY)` returns the next free number, e.g. 3. Redirection and pipes are entirely about *rewiring which fd points to what*.")

h2("2.5  Redirection — making fd 1 point to a file")
para("For `ls > out.txt`, the child opens the file and uses **dup2()** to make stdout (fd 1) refer to it before exec:")
code("""int fd = open("out.txt", O_WRONLY | O_CREAT | O_TRUNC, 0644);
dup2(fd, STDOUT_FILENO);   // fd 1 now points at out.txt
close(fd);                 // the original fd is redundant now
execvp("ls", args);        // ls writes to fd 1 -> lands in out.txt""")
para("`dup2(old, new)` makes `new` refer to the same file as `old`, closing `new` first if it was open.")

h2("2.6  Pipes — connecting two programs")
para("A **pipe** is a one-way in-kernel buffer with two ends: `pipe(fds)` gives `fds[0]` (read end) and `fds[1]` (write end). For `ls | grep .txt` you connect ls's stdout to grep's stdin:")
code("""int fds[2]; pipe(fds);
// In the ls child:    dup2(fds[1], STDOUT_FILENO);  // ls writes into the pipe
// In the grep child:  dup2(fds[0], STDIN_FILENO);   // grep reads from the pipe""")
callout("Classic interview trap:",
        "Why must the PARENT close every pipe fd after forking? Because EOF on a pipe's read end only happens when "
        "ALL write ends are closed. If the parent keeps a write end open, the reader (grep) never sees EOF and hangs "
        "forever — the shell deadlocks. So after forking all children, the parent closes all pipe fds.")

h2("2.7  Signals")
para("A **signal** is an asynchronous notification the kernel delivers to a process — it interrupts whatever the process is doing, runs a handler, then resumes. Signals vsh cares about:")
table(
    ["Signal", "Trigger", "vsh's response"],
    [
        ["SIGINT", "Ctrl+C", "Kill the foreground job, not the shell"],
        ["SIGTSTP", "Ctrl+Z", "Stop (pause) the foreground job"],
        ["SIGCHLD", "A child changed state", "Reap it / update the job table"],
        ["SIGTTIN/SIGTTOU", "Background job touches terminal", "Shell ignores these"],
    ],
    widths=[1.4, 2.2, 2.4],
)
para("Signal handlers must be **async-signal-safe**: do not call malloc/printf, do not block, save and restore `errno`, return fast. (Identical to the rules for an interrupt service routine in embedded code — a connection worth making for Airspan.)")

h2("2.8  Process groups & terminal control")
para("A **process group** is a set of processes sharing a group id (pgid). The terminal delivers Ctrl+C / Ctrl+Z to the *entire foreground process group* at once. The shell:")
bullet("puts each pipeline in its own process group with **setpgid()**, so Ctrl+Z stops every command in a pipeline together;")
bullet("hands the terminal to a group with **tcsetpgrp()** before running a foreground job, and takes it back afterwards;")
bullet("only the foreground group may read the terminal — background readers get SIGTTIN.")

h2("2.9  The environment")
para("The **environment** is a set of `KEY=value` strings every process inherits from its parent — `PATH`, `HOME`, `USER`, etc. vsh stores variables in a hash table (env.c), expands `$VAR` / `${VAR:-default}` / `$?` / `$$`, and passes the right set to each child. `PATH` is how the shell finds `ls` without you typing `/usr/bin/ls`.")

pagebreak()

# ============================================================ 3. RESUME TERMS
h1("3.  Resume Terms Decoded")
para("Your resume bullet reads:")
callout("Resume:",
        "\"Built a POSIX-compatible Linux shell in ~9,600 lines of C from scratch with a recursive-descent parser, "
        "custom line editor (tab completion, Ctrl+R search), and full job control, using zero third-party libraries "
        "(no readline/ncurses)… Engineered an arena allocator and bounds-checked SafeString buffer compiled under "
        "-Wall -Wextra -Werror -Wshadow with ASan/UBSan, driving parser memory leaks to zero as verified by a "
        "215-test unit suite.\"")
para("Every term in it, explained — assume the interviewer points at one and says \"what does this mean?\"")

h3("POSIX")
para("**POSIX** (Portable Operating System Interface) is an IEEE standard that defines the API and behaviour Unix-like systems agree on — system calls like `fork`, `exec`, `pipe`, `dup2`, `waitpid`, plus the standard for what a shell's syntax and builtins should do. Saying vsh is **POSIX-compatible** means it uses only standard POSIX system calls (so it builds on any Unix) and follows POSIX shell grammar (pipes, `&&`, `||`, `if/while/for`, redirections).")
callout("Be honest:",
        "Say \"POSIX-style\" rather than \"fully POSIX-compliant.\" vsh implements most of the grammar but is missing "
        "some POSIX features — command substitution `$(...)`, arithmetic expansion `$((...))`, and `trap`. Knowing "
        "exactly what's missing is more impressive than over-claiming.")

h3("Recursive-descent parser")
para("A parsing technique where each grammar rule becomes one function, and the functions call each other (recurse) to match nested structure. The call hierarchy naturally encodes operator precedence. This gets its own full section (Section 4) because it's the most likely deep-dive question.")

h3("Arena allocator")
para("A memory allocator that grabs one big block up front and hands out pieces by simply bumping a pointer forward (O(1)). Instead of freeing each piece individually, you free the *whole* block in one call. vsh allocates every token and AST node from an arena, then calls `arena_reset()` once per command — that's how it guarantees zero parser leaks. (Detailed in Section 6.)")

h3("SafeString — bounds-checked string buffer")
para("C has no string type — just `char` arrays ending in `'\\0'`, and nothing stops you writing past the end (a **buffer overflow**). SafeString is a growable string that tracks `len` and `cap`, doubles its capacity when full, and checks bounds on every append — so overflow is impossible. It's the same idea as C++ `std::string`, written by hand. (Detailed in Section 6.)")

h3("Custom line editor (no readline)")
para("Most shells link the GNU **readline** library for line editing (arrow keys, history, tab completion). vsh implements all of that itself in ~950 lines using raw terminal mode — so it has zero dependency on readline or ncurses. Demonstrates terminal/UART-style byte-level I/O.")

h3("Job control")
para("The machinery behind Ctrl+Z / `fg` / `bg` / background jobs (`sleep 100 &`): process groups, terminal ownership transfer, and SIGCHLD reaping. (Section 6.)")

h3("Zero third-party libraries")
para("vsh links only the C standard library and POSIX (plus `-lm`, the standard math library, for `calc`). No readline, no ncurses, no anything else. This means it has a tiny, auditable dependency footprint — a real plus for embedded/constrained targets.")

h3("GNU Make")
para("**Make** is a build tool: a `Makefile` lists targets (`release`, `debug`, `test`) and the commands to build them, and rebuilds only what changed. vsh's Makefile defines `make release` (optimised), `make debug`, `make sanitize` (with ASan/UBSan), and `make test`.")

h3("-Wall -Wextra -Werror -Wshadow")
para("Compiler flags that maximise safety:")
bullet("**-Wall** — enable the common warnings.")
bullet("**-Wextra** — enable extra warnings on top.")
bullet("**-Wshadow** — warn when a local variable hides (shadows) an outer one, a common bug source.")
bullet("**-Werror** — treat every warning as a hard error, so the build fails unless the code is warning-clean.")

h3("AddressSanitizer (ASan) / UndefinedBehaviorSanitizer (UBSan)")
para("Compiler-instrumented runtime checkers (`-fsanitize=address,undefined`). **ASan** catches memory errors — buffer overflows, use-after-free, leaks — at the moment they happen, with a stack trace. **UBSan** catches undefined behaviour — signed overflow, misaligned access, bad shifts. You run the test suite under them to prove the code is memory-clean.")

h3("Abstract Syntax Tree (AST)")
para("The tree the parser builds to represent a command's structure. A node is a `struct` tagged with a type (COMMAND, PIPELINE, AND, IF…) plus children. The executor walks this tree to run the command. (Section 6.)")

h3("\"215-test unit suite\"")
para("A custom, dependency-free test framework (`tests/test.h` + `test_main.c`) with 215 assertions across the arena, SafeString, lexer, and parser. `make test` builds and runs them; you run the same suite under ASan to prove zero leaks.")

pagebreak()

# ============================================================ 4. RECURSIVE DESCENT
h1("4.  Recursive-Descent Parsing — In Depth")
para("This is the topic most likely to get a deep follow-up, including \"what are its drawbacks?\" — so know it cold.")

h2("4.1  Why a parser exists at all")
para("The lexer gives a flat list of tokens. A flat list loses structure. Consider:")
code("ls && echo ok || echo fail")
para("Tokens: `[ls] [&&] [echo] [ok] [||] [echo] [fail]`. Two readings are possible:")
bullet("`(ls && echo ok) || echo fail`")
bullet("`ls && (echo ok || echo fail)`")
para("Just as `2 + 3 * 4 = 14` because `*` binds tighter than `+`, shell operators have **precedence**: `|` binds tighter than `&&`/`||`, which bind tighter than `;`. The parser's job is to build a tree that encodes the correct precedence.")

h2("4.2  What recursive descent IS")
para("**Recursive-descent parsing** is a top-down technique where:")
bullet("Each rule in the grammar becomes **one function**.")
bullet("A rule that refers to another rule **calls that rule's function** — hence \"descent.\"")
bullet("Recursive grammar (an expression inside parentheses inside an expression) is handled by **recursion**.")
bullet("Precedence is encoded by **which function calls which** — lower-precedence functions sit higher and call higher-precedence ones.")
para("vsh's (simplified) grammar:")
code("""program  -> list EOF
list     -> pipeline ( ('&&' | '||' | ';' | '&') pipeline )*
pipeline -> command ( '|' command )*
command  -> simple_command | if_stmt | while_stmt | for_stmt | subshell""")
para("Each rule maps to a function: `parse_program`, `parse_list`, `parse_pipeline`, `parse_command`. Here is the real shape of `parse_list`, which handles `&&` / `||`:")
code("""static ASTNode *parse_list(Parser *p) {
    ASTNode *left = parse_pipeline(p);          // parse left side first
    while (check(p, TOK_AND) || check(p, TOK_OR)) {
        TokenType op = cur_token(p)->type;
        advance(p);                             // consume && or ||
        ASTNode *right = parse_pipeline(p);     // parse right side
        ASTNode *node = arena_calloc(p->arena, 1, sizeof(ASTNode));
        node->type  = (op == TOK_AND) ? NODE_AND : NODE_OR;
        node->left  = left;
        node->right = right;
        left = node;        // result becomes the left of the next iteration
    }
    return left;
}""")
para("Because `parse_list` calls `parse_pipeline` (which calls `parse_command`), the tree guarantees that `|` binds tighter than `&&`/`||` — exactly right, with no precedence table needed.")

h2("4.3  The resulting AST")
para("For `ls | grep .txt && echo done`:")
code("""          AND
        /      \\
   PIPELINE     COMMAND(echo done)
   /     \\
 CMD(ls)  CMD(grep .txt)""")
para("Every node is `arena_calloc()`'d — allocated from the arena and zeroed. The whole tree, however deep, is freed by the next `arena_reset()`.")

h2("4.4  Why vsh uses it")
bullet("**Readable & maintainable** — the code reads exactly like the grammar; you can hand-trace it.")
bullet("**No external tools** — no yacc/bison/ANTLR, no generated code; fits the \"zero dependencies\" goal.")
bullet("**Great error messages** — you control exactly where and how errors are reported.")
bullet("**Industry-standard** — GCC, Clang, and most production compilers use hand-written recursive-descent parsers for exactly these reasons.")

h2("4.5  Drawbacks of recursive descent  (know these!)")
para("Interviewers love this question because over-claiming is easy. The honest limitations:")
table(
    ["Drawback", "Explanation"],
    [
        ["Cannot handle left-recursion",
         "A rule like `expr -> expr '+' term` calls itself with no token consumed first -> infinite recursion / stack overflow. You must rewrite the grammar into a loop (the `while` in parse_list) or right-recursion."],
        ["Precedence is hand-coded",
         "Each precedence level needs its own function/loop. For a language with many operator levels this is verbose and error-prone — a table-driven or Pratt parser scales better."],
        ["Limited lookahead (LL)",
         "It's an LL parser: it decides which rule to use from the next token(s). Grammars needing long lookahead, or that are ambiguous, can't be parsed directly without hacks or backtracking."],
        ["Backtracking can be exponential",
         "If you add naive backtracking to resolve ambiguity, worst-case time can blow up. vsh avoids this by keeping the grammar predictive (LL(1)-ish)."],
        ["Stack overflow on deep nesting",
         "Recursion depth = nesting depth. Deeply nested input (thousands of parentheses) can overflow the call stack, since there's no explicit stack to bound."],
        ["Less powerful than LR",
         "LR/LALR parsers (yacc/bison) accept a strictly larger class of grammars and catch grammar conflicts at build time. Recursive descent can't express some grammars without rewriting."],
        ["Manual & repetitive",
         "You write and maintain every function by hand; a parser generator would derive it from the grammar automatically."],
    ],
    widths=[2.0, 4.3],
)
callout("Interview answer (drawbacks):",
        "\"Recursive descent can't handle left-recursive grammar rules — they'd recurse infinitely — so I rewrite "
        "those as loops. Precedence is encoded by hand in the call hierarchy, which is fine for a shell's handful of "
        "operators but doesn't scale to a language with many levels; there a Pratt or table-driven parser is cleaner. "
        "It's LL with limited lookahead, so it can't parse ambiguous grammars that an LR parser like bison could, and "
        "deeply nested input risks call-stack overflow because recursion depth equals nesting depth. I chose it anyway "
        "because the grammar is small, it needs no external tools, and the code reads like the grammar — the same "
        "trade-off GCC and Clang make.\"")

h2("4.6  Alternatives (good to name-drop)")
bullet("**Pratt parsing (top-down operator precedence)** — recursive descent + a precedence table; elegant for expression-heavy languages.")
bullet("**LR / LALR (yacc, bison)** — bottom-up, table-driven, handles a larger grammar class, detects conflicts at build time; but generated code is harder to debug.")
bullet("**PEG / packrat** — formalises ordered-choice with memoised backtracking for linear time.")

callout("Bonus point:",
        "vsh actually contains TWO recursive-descent parsers. Besides the shell grammar, the `calc` builtin uses a "
        "second one (parse_expr -> parse_term -> parse_factor) to evaluate math like `2 + 3 * sin(pi/2)` — a textbook "
        "expression grammar where each function is one precedence level. Great to mention as proof you understand the "
        "pattern, not just copied it once.")

pagebreak()

# ============================================================ 5. SYSTEMS C REFRESHER
h1("5.  Systems-C Refresher (beyond the basics)")
para("You know basic C. These are the specifically systems-level pieces vsh relies on that a typical intro course skips.")

h3("Strings are just char arrays")
para("There's no string type. `char *s = \"Jacob\"` is a pointer to `'J'`; the run ends at a `'\\0'` byte. `strlen` counts until `'\\0'`. Writing past the buffer corrupts neighbouring memory — the reason SafeString exists.")

h3("malloc / free / realloc")
para("`malloc(n)` returns heap memory (or NULL — always check). `free(p)` releases it; forgetting is a **leak**, doing it twice is a crash, using after it is a **dangling pointer**. `realloc` grows a block (SafeString's doubling uses it). In C, `malloc` returns `void *` and needs no cast.")

h3("struct and ->")
para("C has no classes — a `struct` just groups data. Pass a pointer (`Shell *shell`) not a copy, and access members with `->`. Every vsh function takes `Shell *shell` so they all share one state object.")

h3("Function pointers")
para("A variable holding the address of a function: `int (*fn)(Shell*, char**)`. The builtin table is an array of `{name, function pointer}` — dispatch is a loop comparing names, then calling through the pointer. This is C's version of a virtual method table.")

h3("write() vs printf()")
para("`printf` is buffered and uses global state — unsafe in a signal handler. `write(fd, buf, n)` is a direct, unbuffered system call — safe in handlers and used throughout the line editor.")

h3("va_list — variadic functions")
para("The `...` in a function signature (like `printf`) means \"any number of extra arguments.\" `va_start`/`va_arg`/`va_end` walk them. A `va_list` can't be reused once read — that's why SafeString's `appendf` runs two separate passes (measure, then write).")

h3("The double-pass vsnprintf trick")
para("`vsnprintf(NULL, 0, fmt, ap)` writes nothing but returns how many bytes the result *would* take. Measure first, allocate exactly that, then write — overflow becomes impossible. Used in SafeString and the prompt builder.")

h3("C99 flexible array member")
para("`char data[];` as the last struct member has no size; you allocate `sizeof(struct) + N` and the N bytes sit right after the struct — one allocation, not two. The arena's pages use this.")

pagebreak()

# ============================================================ 6. ARCHITECTURE + FILES
h1("6.  Architecture & File-by-File Walkthrough")

h2("6.1  The big picture")
code("""You type a command
        |
        v
+-----------------------------------------------+
| vsh_readline.c  - Line editor (raw mode)      |
|   arrows, history, tab-complete, Ctrl+R       |
|   returns: char *line                         |
+-----------------------+-----------------------+
                        v
+-----------------------------------------------+
| lexer.c  - Tokenizer (finite state machine)   |
|   returns: TokenList  [WORD][PIPE][WORD]...    |
+-----------------------+-----------------------+
                        v
+-----------------------------------------------+
| parser.c - Recursive-descent parser           |
|   returns: ASTNode *  (a tree)                |
+-----------------------+-----------------------+
                        v
+-----------------------------------------------+
| executor.c - AST walker                        |
|   +-- pipeline.c     fork + pipe + exec        |
|   +-- job_control.c  process groups, fg/bg     |
|   +-- builtins.c     cd, export, exit, ...      |
|   +-- env.c          $VAR expansion             |
+-----------------------------------------------+

Support libs used everywhere:  arena.c (allocator)
                               safe_string.c (string builder)
                               wildcard.c (glob)""")

h2("6.2  Data flow for one command")
para("For `ls -la | grep .txt`:")
numbered("`vsh_readline()` returns the line string.")
numbered("`arena_reset()` clears the previous command's tokens/AST.")
numbered("`lexer_tokenize()` -> `[WORD ls][WORD -la][PIPE][WORD grep][WORD .txt][EOF]`.")
numbered("`parser_parse()` -> PIPELINE node with two COMMAND children.")
numbered("Executor dispatches to pipeline.c, which makes 1 pipe, forks 2 children, wires them, waits.")
numbered("grep's output reaches the terminal.")
numbered("`arena_reset()` frees every token and node in O(1).")

# ---- the modules ----
modules = [
    ("arena.c / arena.h — the memory manager",
     [("Job", "Pre-allocate a big page; hand out memory by bumping a pointer; free everything at once with arena_reset()."),
      ("Why", "The parser builds dozens of nodes per command. Individual free() calls invite leaks (forget one) and crashes (free twice). The arena makes both impossible."),
      ("Key structs", "An ArenaPage is a linked-list node holding `size`, `used`, and a C99 flexible array `char data[]`. The Arena tracks `head`, `current`, and `page_size` (usually 4096)."),
      ("Bump allocation", "arena_alloc rounds the size up to a multiple of 8 (alignment), checks it fits in the current page, returns `data + used`, and adds to `used`. If the page is full, it links a new page. Always O(1)."),
      ("align_up trick", "`(size + 7) & ~7` rounds up to the next multiple of 8 so every allocation is CPU-aligned."),
      ("Reset", "arena_reset frees every page after the first, sets the first page's `used` back to 0, and reuses it. One call frees the whole parse tree; the first page is kept warm for the next command."),
      ("Interview", "\"A page-based bump allocator: O(1) allocation, O(1) reset. Every token and AST node lives in the arena, so one arena_reset() per command frees the entire parse state — parser leaks are structurally impossible. Same idea as an RTOS memory pool.\"")]),

    ("safe_string.c / safe_string.h — the string builder",
     [("Job", "A growable, bounds-checked string for assembling tokens, prompts, and output without overflow."),
      ("Struct", "`{ char *data; size_t len; size_t cap; }` — current length and total capacity tracked separately."),
      ("2x growth", "When an append won't fit, capacity doubles (`realloc`). Doubling gives amortised O(1) appends — N appends cost O(N), not O(N^2). Same strategy as std::string / std::vector."),
      ("appendf double-pass", "vsnprintf(NULL,0,...) measures the needed size, then it grows, then it writes — format overflow is impossible. Two va_start/va_end blocks because a va_list can't be reused."),
      ("Interview", "\"A bounds-checked dynamic string with 2x growth for amortised O(1) appends, and a measure-then-write vsnprintf pattern so formatted output can't overflow.\"")]),

    ("lexer.c / lexer.h — the tokenizer",
     [("Job", "Turn the raw line into a typed token list: WORD, PIPE, AND (&&), OR (||), SEMICOLON, AMPERSAND, REDIRECT_IN/OUT, APPEND (>>), the keywords if/then/else/fi/while/do/done/for/in, and EOF."),
      ("Why not split on spaces", "Quotes (`echo \"a b\"` is one word), escapes (`a\\ b`), and operators with no surrounding spaces (`grep foo|bar`) all break naive splitting. You need a state machine."),
      ("How", "It walks the input one char at a time tracking quoting state (unquoted / single / double). Inside single quotes everything is literal; inside double quotes spaces/operators are literal but `$` still expands later."),
      ("Greedy matching", "For 1-or-2-char operators it peeks ahead: see `>`, peek — if the next is `>` emit APPEND (`>>`), else emit REDIRECT_OUT. Always test the longer operator first (`||` before `|`)."),
      ("Memory", "Each token's string is arena_strdup'd; the whole TokenList is arena-allocated and vanishes on the next reset."),
      ("Interview", "\"A hand-written finite state machine tracking quoting state, building tokens in a SafeString, using greedy lookahead for multi-char operators. The same FSM pattern shows up in AT-command and protocol-frame parsers.\"")]),

    ("parser.c / parser.h — recursive-descent parser",
     [("Job", "Turn the flat token list into an AST that encodes precedence (see Section 4 for the deep dive)."),
      ("Shape", "One function per grammar rule: parse_program -> parse_list -> parse_pipeline -> parse_command, plus parse_if/while/for. Helpers check()/advance()/expect() manage the token cursor."),
      ("ASTNode", "A tagged struct: `type` (NODE_COMMAND, NODE_PIPELINE, NODE_AND, NODE_IF, ...), plus `args/argc/redirs` for commands, `left/right` for binary nodes, `condition/then/else` for if, `body` for loops/functions. Each is arena_calloc'd (zeroed)."),
      ("Interview", "\"Recursive descent — each rule is a function, the call hierarchy encodes precedence (pipelines bind tighter than &&/||), and every node is arena-allocated so the whole tree frees in one reset.\"")]),

    ("executor.c / executor.h — the AST walker",
     [("Job", "Walk the AST and actually run it. A switch on node type dispatches to exec_command / exec_pipeline / exec_and / exec_or / exec_sequence / exec_if / exec_while / exec_for / exec_background."),
      ("Single command", "Check the builtin table first (no fork). Otherwise fork; in the child: setpgid, reset signals to default, apply redirections, execvp; in the parent: register the job and wait."),
      ("Why builtins don't fork", "`cd` calls chdir(), which only changes the calling process's directory. If cd ran in a child, the child would change dir and exit — the shell wouldn't move. So cd/export/unset/exit/source/jobs run in-process."),
      ("&& / || short-circuit", "exec_and runs the right side only if the left returned 0 (success); exec_or only if the left was non-zero. Uses shell->last_status."),
      ("Expansion", "Before exec, arguments go through variable expansion ($VAR), tilde (~), and glob/wildcard matching (wildcard.c)."),
      ("Known bug", "exec_function serialises the function body's AST pointer into an env var as a hex string. But arena_reset() frees that AST before the function is ever called -> dangling pointer -> segfault. Fix: store bodies in a persistent malloc'd hash map outside the parse arena. (See Section 8.)"),
      ("Interview", "\"An AST walker; simple commands fork/exec, builtins run in-process because they mutate shell state, &&/|| short-circuit on exit status, and redirections are applied in the child before exec via dup2.\"")]),

    ("pipeline.c / pipeline.h — connecting programs",
     [("Job", "Run `a | b | c`: create N-1 pipes for N commands, fork N children, dup2 each child's stdin/stdout to the right pipe ends, close all pipe fds in the parent, wait for all."),
      ("Middle command", "Child i reads from pipe[i-1] and writes to pipe[i]; first reads the terminal, last writes the terminal."),
      ("The critical close", "After forking, the parent MUST close every pipe fd. Otherwise a downstream reader never gets EOF (a write end is still open in the parent) and the shell deadlocks. This is the #1 pipe interview question."),
      ("Interview", "\"N-1 pipes for N commands, dup2 to wire each stage, and — the detail people miss — the parent closes all pipe fds after forking so EOF propagates and nothing hangs.\"")]),

    ("job_control.c / job_control.h — fg / bg / Ctrl+Z",
     [("Job", "Implement the Unix job-control model: background jobs (&), fg, bg, jobs, stop/continue with Ctrl+Z."),
      ("Process groups", "Each pipeline gets its own group via setpgid so terminal signals hit the whole group. tcsetpgrp transfers terminal ownership to the foreground group and back to the shell afterwards."),
      ("Init", "The shell puts itself in its own group, grabs the terminal, and ignores SIGTSTP/SIGTTIN/SIGTTOU so Ctrl+Z never freezes the shell itself."),
      ("SIGCHLD handler", "Installed with SA_RESTART | SA_NOCLDSTOP. It loops waitpid(-1, ..., WNOHANG|WUNTRACED|WCONTINUED) to reap every child that changed state, and saves/restores errno. WNOHANG = never block in a handler; SA_RESTART = auto-retry interrupted read() so the line editor doesn't crash when a background job exits."),
      ("Interview", "\"Process groups + tcsetpgrp for terminal ownership; a SIGCHLD handler that loops with WNOHANG, saves/restores errno, and uses SA_RESTART. The handler discipline is identical to an ISR — no blocking, save volatile state, return fast.\"")]),

    ("vsh_readline.c / vsh_readline.h — the custom line editor",
     [("Job", "Replace the readline library: read keystrokes live, support editing, history, search, tab completion — in ~950 lines."),
      ("Raw mode", "tcsetattr clears ECHO and ICANON so every keypress arrives immediately (instead of the OS buffering a whole line). The editor echoes characters itself and restores the old terminal settings on exit."),
      ("ANSI sequences", "Arrow keys arrive as 3 bytes (ESC [ A/B/C/D); the editor parses them. Screen updates are done by writing ANSI codes (move cursor, erase line) with write(), not printf() (signal-safety + no buffering)."),
      ("Features", "Up/Down history, Ctrl+R reverse search, Ctrl+A/E home/end, Ctrl+K/U/Y kill-and-yank, Ctrl+W delete-word, Tab completion (scans PATH for commands and the cwd for files, shows a columnar menu)."),
      ("Known bug", "prompt_len = strlen(prompt) miscounts when the prompt has ANSI colour codes — those bytes are invisible on screen, so cursor math is off for coloured prompts. Fix: measure visible width, skipping escape sequences."),
      ("Interview", "\"Raw mode via tcsetattr, manual ANSI parsing for special keys, redraw with write() for signal-safety. It's effectively a tiny UART-style byte protocol handler — the same skills as a serial driver.\"")]),

    ("env.c / env.h — variables & expansion",
     [("Job", "Store shell/environment variables in a hash table and perform expansion: $VAR, ${VAR:-default}, $? (last status), $$ (pid), $#, $@, plus export/unset and exporting to children's environ."),
      ("Interview", "\"A hash-table variable store plus an expansion engine the executor runs over each argument before exec.\"")]),

    ("history.c / history.h — command history",
     [("Job", "A ring buffer of past commands with persistence to a history file, used by Up/Down, Ctrl+R, and history expansion (!!, !N, !-N, !prefix)."),
      ("Interview", "\"A ring buffer persisted to disk; the line editor and the !-expansion both read from it.\"")]),

    ("wildcard.c / wildcard.h — globbing",
     [("Job", "Expand glob patterns (*, ?, [...]) against the filesystem and do tilde (~) expansion, so `ls *.c` becomes the matching filenames before exec."),
      ("Interview", "\"A hand-written glob matcher invoked during argument expansion.\"")]),

    ("shell.c / shell.h — the REPL & orchestration",
     [("Job", "The heart: the read-eval-print loop, prompt rendering (two-line powerline prompt with time, user@host, shortened cwd, git branch, colour-coded exit status), signal setup, and sourcing ~/.vshrc on interactive startup. shell.h defines the central Shell state struct passed everywhere."),
      ("Interview", "\"shell.c owns the REPL and the Shell state object; every subsystem hangs off the one Shell* that's threaded through the code.\"")]),

    ("main.c — entry point",
     [("Job", "Parse CLI options and choose a mode: interactive (`./vsh`), single command (`./vsh -c \"...\"`), or run a script file (`./vsh script.sh`). Initialises the shell, runs it, cleans up."),
      ("Interview", "\"Thin entry point that selects interactive vs -c vs script mode and hands off to shell.c.\"")]),

    ("builtins.c + builtins/ — the builtin commands",
     [("Job", "A registry (name -> function pointer) and 16 builtin implementations. Standard ones: cd, pwd, echo, export, unset, alias/unalias, history, source/'.', type, jobs/fg/bg, pushd/popd/dirs, exit, help, return, local."),
      ("Why a table", "Dispatch is a loop matching the command name against the table, then a call through the function pointer — C's idiom for polymorphic dispatch."),
      ("Showcase builtins", "calc, httpfetch, sysinfo, watch, colors — covered in Section 7."),
      ("Interview", "\"A function-pointer table; the executor checks it before forking so state-changing commands run in-process.\"")]),
]

for title_text, items in modules:
    h2(title_text)
    for label, body in items:
        if label == "Interview":
            callout("Interview:", body.strip('"'))
        else:
            p = doc.add_paragraph()
            r = p.add_run(label + ": ")
            r.bold = True
            r.font.color.rgb = DARK
            add_runs(p, body)

h2("6.x  Headers, tests & the build")
para("**include/ (13 headers)** — the public interface of each module: types and function prototypes. Separating declarations (.h) from definitions (.c) is how C does modularity and lets files compile independently.")
para("**tests/ (215 tests)** — a tiny custom framework (`test.h` macros, `test_main.c` runner) covering arena, SafeString, lexer, and parser. No external test library, matching the zero-dependency goal. Run under ASan to prove zero leaks.")
para("**Makefile** — `make release` (optimised `./vsh`), `make debug`, `make sanitize` (ASan+UBSan), `make test`, `make clean`, `make install`. Compiles everything under `-Wall -Wextra -Werror -Wshadow`.")

pagebreak()

# ============================================================ 7. NEW / SHOWCASE FEATURES
h1("7.  The \"New\" / Showcase Features")
para("Beyond a textbook shell, vsh has distinctive features worth highlighting — they show range (sockets, parsing, terminal graphics) and give you concrete things to talk about.")

h2("7.1  calc — a second recursive-descent parser")
para("A built-in math evaluator: `calc 2 + 3 * sin(pi/2)`. It has its own lexer + recursive-descent parser (`parse_expr -> parse_term -> parse_factor`), supporting `+ - * / ^`, parentheses, functions (`sin, cos, tan, sqrt, log, exp, abs`), and constants (`pi, e`). This is the cleanest textbook example of the parsing technique from Section 4 — each function is exactly one precedence level. **This is your strongest \"I understand parsing\" talking point.**")

h2("7.2  httpfetch — raw-socket HTTP")
para("`httpfetch <url>` does an HTTP GET using **raw BSD sockets** — `getaddrinfo`, `socket`, `connect`, `send`, `recv` — with no curl or HTTP library. It writes the request line and headers by hand, reads the response, and follows redirects. Demonstrates network programming at the syscall level (directly relevant to a telecom/networking company like Airspan).")

h2("7.3  sysinfo — a system dashboard")
para("A colour dashboard reading OS, kernel, CPU, memory, disk, and uptime — by parsing `/proc` and `/sys` and calling `uname`/`sysconf`. Shows you can pull structured data out of the Linux kernel's pseudo-filesystems.")

h2("7.4  watch — periodic re-execution")
para("`watch -n 2 date` re-runs a command every N seconds, clearing the screen each time using ANSI codes — implemented with the shell's own executor plus a timed loop.")

h2("7.5  colors — terminal graphics")
para("Prints the 256-colour palette and a true-colour (24-bit) gradient using ANSI escape codes — proof you understand terminal control sequences (the same machinery as the line editor).")

h2("7.6  The powerline prompt with git detection")
para("A two-line prompt showing time, user@host, a shortened working directory, a colour-coded exit-status indicator, and the current **git branch** — found by walking up the directory tree looking for `.git/HEAD` and reading it. No libgit2; just file I/O.")

h2("7.7  Shell-language features")
bullet("**Control flow** — if/elif/else/fi, while/do/done, for/in/do/done, subshells, block grouping.")
bullet("**Variable expansion** — $VAR, ${VAR:-default}, $?, $$, $#, $@.")
bullet("**History expansion** — !!, !N, !-N, !prefix.")
bullet("**Aliases** with recursive-expansion detection.")
bullet("**Redirections** including fd targeting like `2>&1`.")
bullet("**~/.vshrc** sourced on interactive startup (your own rc file, like .bashrc).")

callout("Framing for the interview:",
        "\"calc and httpfetch are the two I'd demo: calc reuses the recursive-descent technique on a math grammar, "
        "and httpfetch does HTTP over raw sockets with no library — both show I understand the layer underneath the "
        "abstractions most people just import.\"")

pagebreak()

# ============================================================ 8. BUGS / AUDIT
h1("8.  Known Bugs & Honest Resume Audit")
para("Being able to name your project's limitations precisely is a senior-level signal. Don't hide these — own them.")

h2("8.1  Honest audit of the resume claims")
table(
    ["Claim", "Verdict", "Honest detail"],
    [
        ["C from scratch, zero external deps", "True", "Only -lm (standard math lib). Verified in the Makefile."],
        ["Custom line editor, no readline/ncurses", "True", "~950 lines of raw terminal I/O."],
        ["Arena allocator, zero-leak parsing", "True", "Genuinely solid: O(1) alloc and reset."],
        ["Bounds-checked SafeString", "True", "Proper 2x growth, double-pass vsnprintf."],
        ["Full job control (signals, fg/bg)", "True", "Real SIGCHLD, process groups, tcsetpgrp, SA_RESTART."],
        ["Pipelines, redirection, AST dispatch", "Mostly true", "Heredoc (<<) is parsed but the executor stubs it as 'not yet implemented'."],
        ["Memory-safe", "Mostly true", "exec_function has a real dangling-pointer bug (below)."],
        ["POSIX-compatible", "Stretch", "Say 'POSIX-style'. Missing: $() command substitution, $(()) arithmetic, trap."],
    ],
    widths=[2.2, 1.0, 3.1],
)

h2("8.2  Bug 1 — exec_function dangling pointer")
para("Shell functions (`f() { ls; }`) are broken. The body is an arena-allocated AST node; the code stores its pointer as a hex string in an env var and reads it back when the function is called. But `arena_reset()` runs after every command, freeing that AST first — so the stored pointer is dangling and calling the function segfaults.")
callout("How to say it:",
        "\"Functions have a known dangling-pointer bug: I serialised the body's AST pointer into an env var, but the "
        "arena that owns that AST is reset after every command, so the pointer is invalid by the time the function "
        "runs. The fix is a persistent hash map holding malloc'd copies of function bodies, outside the parse arena.\"")

h2("8.3  Bug 2 — heredoc stubbed")
para("`<< EOF ... EOF` is lexed and parsed (the token type exists) but the executor prints \"not yet implemented.\" ~50 lines of work using a temp file or memfd to feed the body to stdin.")

h2("8.4  Bug 3 — coloured-prompt cursor math")
para("prompt_len uses strlen, which counts invisible ANSI colour bytes, so cursor positioning is wrong for coloured prompts. Fix: compute visible width, skipping escape sequences.")

callout("Meta-point:",
        "If asked \"what would you fix first?\" -> the exec_function bug (functions are completely broken), then "
        "heredoc, then the prompt-width calculation. Naming a prioritised fix list is exactly what they want to hear.")

pagebreak()

# ============================================================ 9. Q&A BANK
h1("9.  Interview Q&A Bank")

qa = [
    ("Memory", [
        ("What is a memory leak, and how does vsh avoid them in the parser?",
         "Allocated heap memory that's never freed because the pointer is lost. vsh arena-allocates all tokens and AST nodes; one arena_reset() per command frees everything, so nothing can be forgotten."),
        ("What is a dangling pointer?",
         "A pointer to already-freed memory; dereferencing it is undefined behaviour. The exec_function bug is exactly this — the AST is reset before the saved pointer is used."),
        ("What is a buffer overflow?",
         "Writing past the end of a fixed buffer, corrupting adjacent memory. C won't stop you. SafeString prevents it by checking capacity before every write."),
        ("Stack vs heap?",
         "Stack memory is auto-managed per function call (fast, small ~8MB); heap (malloc/free) has a lifetime you control (large, manual). vsh's tokens/nodes go on the heap via the arena because they outlive the functions that build them."),
        ("What is memory alignment and why does the arena care?",
         "Data starting at an address divisible by its size; CPUs access aligned data faster (or require it). arena align_up rounds every allocation to a multiple of 8."),
    ]),
    ("OS / Linux", [
        ("What does fork() return?",
         "Twice: the child's PID in the parent, 0 in the child, -1 on failure (to the caller only)."),
        ("Difference between fork and exec?",
         "fork copies the process; exec replaces it with a new program. Run a program by forking then exec-ing in the child so the parent survives."),
        ("How does a pipe + dup2 work?",
         "A pipe is a one-way kernel buffer with a read and write fd. dup2(old,new) makes new refer to old's file. For a|b: b's stdin = pipe read end, a's stdout = pipe write end."),
        ("Why must the parent close pipe ends?",
         "EOF on the read end needs ALL write ends closed. If the parent keeps one open, the reader hangs forever -> deadlock."),
        ("What is SIGCHLD and how do you handle it safely?",
         "Sent when a child changes state. Handle with a loop of waitpid(WNOHANG) (never block), save/restore errno, install with SA_RESTART so interrupted syscalls retry."),
        ("Why save/restore errno in a signal handler?",
         "errno is global; the handler can interrupt code mid-error-check, and syscalls in the handler overwrite it. Save on entry, restore on exit."),
        ("What is a process group and why use one?",
         "A set of processes sharing a pgid; the terminal signals the whole foreground group. The shell puts each pipeline in its own group so Ctrl+Z stops it as a unit."),
        ("What does tcsetpgrp do?",
         "Sets the foreground process group for a terminal — i.e. hands the terminal to a job before waiting on it, and takes it back after."),
        ("Canonical vs raw terminal mode?",
         "Canonical (cooked) mode buffers a whole line and handles editing for you; raw mode (ECHO/ICANON off) delivers each keypress immediately — required for a custom line editor."),
    ]),
    ("Parsing", [
        ("What is recursive descent and how does it encode precedence?",
         "Top-down parsing with one function per grammar rule; the call hierarchy encodes precedence (parse_list calls parse_pipeline, so | binds tighter than &&/||)."),
        ("What are its drawbacks?",
         "Can't do left-recursion (infinite loop), precedence is hand-coded, limited lookahead (LL), risk of stack overflow on deep nesting, and less powerful than LR/bison. See Section 4.5."),
        ("Lexer vs parser?",
         "The lexer groups characters into typed tokens (no structure); the parser groups tokens into a tree (structure + precedence)."),
        ("What's an AST?",
         "A tagged tree representing the command's structure; the executor walks it. Nodes like COMMAND, PIPELINE, AND, IF."),
    ]),
    ("The project", [
        ("Walk me through `ls | wc -l`.",
         "readline reads the line; arena_reset clears state; lexer -> [WORD ls][PIPE][WORD wc][WORD -l][EOF]; parser -> PIPELINE with two COMMANDs; pipeline.c makes 1 pipe, forks ls (stdout->pipe), forks wc (stdin->pipe), closes pipe fds in the parent, waits both; arena_reset frees everything."),
        ("Why is cd a builtin?",
         "chdir only changes the calling process's directory. In a forked child it'd change then exit, leaving the shell put. So cd runs in-process — like export, unset, exit."),
        ("What would you improve first?",
         "Fix the exec_function dangling pointer (persistent hash map for function bodies), then implement heredoc, then the coloured-prompt width calc."),
        ("Why no readline / external libs?",
         "To learn the layer underneath and to keep the dependency footprint tiny and auditable — valuable for embedded/constrained targets."),
    ]),
]

for cat, pairs in qa:
    h2(cat)
    for q, a in pairs:
        p = doc.add_paragraph()
        r = p.add_run("Q:  ")
        r.bold = True
        r.font.color.rgb = ACCENT
        add_runs(p, q)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.25)
        r = p2.add_run("A:  ")
        r.bold = True
        r.font.color.rgb = DARK
        add_runs(p2, a)

pagebreak()

# ============================================================ 10. AIRSPAN CONNECTION
h1("10.  Why This Matters for Airspan (the embedded/telecom bridge)")
para("Airspan builds 4G/5G radio network equipment — embedded Linux, C, real-time constraints, networking. Map every vsh skill onto something they care about:")
table(
    ["vsh concept", "Where in vsh", "Airspan-relevant parallel"],
    [
        ["Manual memory management", "arena.c, safe_string.c", "Embedded has no garbage collector"],
        ["Arena / memory pool", "arena.c", "RTOS memory pools (FreeRTOS, Zephyr) — no heap fragmentation"],
        ["fork/exec/wait lifecycle", "executor.c, pipeline.c", "Process management in embedded Linux DU/CU stacks"],
        ["Async-signal-safe handlers", "job_control.c", "Interrupt service routines — no blocking, save state, return fast"],
        ["File descriptors / dup2", "executor.c, pipeline.c", "Unix sockets & network fds in a 5G stack"],
        ["Raw terminal byte I/O", "vsh_readline.c", "UART / serial drivers"],
        ["Finite state machine", "lexer.c", "Protocol frame decoders, AT-command parsers"],
        ["Recursive-descent parsing", "parser.c, calc", "Config / message parsing (e.g. O-RAN, ASN.1-style)"],
        ["Raw sockets", "httpfetch", "Network programming at the syscall level"],
        ["Pipe EOF / backpressure", "pipeline.c", "DMA ring buffers — consumer must not block with no producer"],
    ],
    widths=[1.9, 1.7, 2.7],
)
callout("One-liner to close with:",
        "\"The reason I value this project is that it forced me to work one layer below the abstractions — manual "
        "memory, raw fds, signal-safe code, byte-level terminal I/O. That's exactly the layer embedded and telecom "
        "systems live at, which is why I'm excited about the systems/software role at Airspan.\"")

pagebreak()

# ============================================================ 11. 4-MIN PITCH
h1("11.  The 4-Minute Deep-Dive Pitch")
para("When they say \"walk me through it in detail.\" Practise this out loud and time it.", italic=True)
para("\"vsh is a POSIX-style Unix shell I wrote from scratch in C — about 9,600 lines, zero third-party libraries, not even readline.")
para("The core is a four-stage pipeline. A custom line editor reads raw terminal keystrokes — I put the terminal in raw mode with tcsetattr, parse arrow keys as ANSI escape sequences, and redraw with write() so it's signal-safe. A hand-written lexer tokenises the line — it's a finite state machine that tracks quoting and does greedy matching for multi-character operators. A recursive-descent parser builds an abstract syntax tree, where the call hierarchy encodes precedence so pipelines bind tighter than && and ||. Then an executor walks the tree — forking processes, wiring pipes with dup2, and running builtins in-process.")
para("The part I find most interesting is memory. Every token and AST node is allocated from an arena — a page-based bump allocator. Allocation is an O(1) pointer bump; after each command, one arena_reset() frees the whole parse tree in O(1). There are no individual frees, so parser leaks are structurally impossible. It's the same pattern as an RTOS memory pool.")
para("For job control I implemented the full Unix model: setpgid for process groups, tcsetpgrp to hand the terminal to a job, and a SIGCHLD handler installed with SA_RESTART that loops with WNOHANG to reap children without blocking — saving and restoring errno, exactly like an interrupt handler.")
para("The honest limitation is that shell functions have a dangling-pointer bug: I serialised the body's AST pointer into an environment variable, but the arena that owns it is reset every command, so it's invalid by call time. The fix is a persistent hash map for function bodies outside the parse arena. Heredocs are parsed but not yet executed, and I'd call it POSIX-style rather than fully compliant since command substitution isn't there yet.")
para("It builds under -Wall -Wextra -Werror with AddressSanitizer and UBSan, and a 215-test suite verifies the arena, the string buffer, the lexer, and the parser with zero leaks.")
para("For Airspan specifically, the arena maps onto embedded memory pools, the signal discipline is identical to ISR rules, and the pipe and process-lifecycle work is what embedded Linux does constantly at the process-coordination layer.\"")

hr()
para("End of guide — read sections 1–4 first, then drill the modules you're most likely to be quizzed on (arena, parser, executor, pipeline, job_control). Good luck, Jacob.", italic=True, color=GREY)

# ---------------------------------------------------------------- save
out = "/home/jake/Documents/Projects/Airspan-Interview-Prep/vsh-study-guide.docx"
doc.save(out)
print("Saved:", out)
